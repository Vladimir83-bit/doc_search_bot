from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import os
import re

class DocumentParser:
    """Класс для извлечения текста из разных форматов документов"""
    
    @staticmethod
    def parse_text(file_path):
        """Чтение обычных текстовых файлов (.txt)"""
        try:
            encodings = ['utf-8', 'cp1251', 'windows-1251']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            print(f"Ошибка чтения TXT: {e}")
            return ""

    @staticmethod
    def parse_pdf(file_path):
        """Извлечение текста из PDF"""
        text = ""
        try:
            reader = PdfReader(file_path)
            max_pages = 50
            for i, page in enumerate(reader.pages):
                if i >= max_pages:
                    break
                text += page.extract_text() or ""
        except Exception as e:
            print(f"Ошибка чтения PDF: {e}")
        return text

    @staticmethod
    def parse_docx(file_path):
        """Чтение документов Word"""
        try:
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Ошибка чтения DOCX: {e}")
            return ""

    @staticmethod
    def parse_excel(file_path):
        """Обработка Excel файлов"""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            print(f"Ошибка чтения Excel: {e}")
            return ""

    @classmethod
    def parse_file(cls, file_path):
        """Автоматическое определение типа файла и его обработка"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return cls.parse_text(file_path)
        elif ext == '.pdf':
            return cls.parse_pdf(file_path)
        elif ext == '.docx':
            return cls.parse_docx(file_path)
        elif ext in ('.xlsx', '.xls'):
            return cls.parse_excel(file_path)
        else:
            print(f"Неподдерживаемый формат: {ext}")
            return ""

    @staticmethod
    def find_text_with_context(file_path, search_text, context_words=20):
        """Находит текст и возвращает абзацы/фрагменты с контекстом"""
        try:
            content = DocumentParser.parse_file(file_path)
            if not content:
                return []
            
            search_text_lower = search_text.lower()
            content_lower = content.lower()
            
            # Разбиваем текст на абзацы/предложения
            paragraphs = []
            if file_path.endswith('.docx'):
                # Для Word документов используем настоящие абзацы
                doc = Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            else:
                # Для других форматов разбиваем по переносам строк
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            
            # Если не получилось разбить на абзацы, разбиваем по точкам
            if not paragraphs:
                paragraphs = [s.strip() for s in content.split('.') if s.strip()]
            
            found_paragraphs = []
            
            # Ищем в каждом абзаце
            for paragraph in paragraphs:
                if search_text_lower in paragraph.lower():
                    # Ограничиваем длину абзаца для удобства чтения
                    if len(paragraph) > 500:
                        # Находим позицию искомого текста
                        pos = paragraph.lower().find(search_text_lower)
                        start = max(0, pos - 100)
                        end = min(len(paragraph), pos + len(search_text) + 100)
                        shortened_paragraph = paragraph[start:end]
                        if start > 0:
                            shortened_paragraph = "..." + shortened_paragraph
                        if end < len(paragraph):
                            shortened_paragraph = shortened_paragraph + "..."
                        found_paragraphs.append(shortened_paragraph)
                    else:
                        found_paragraphs.append(paragraph)
            
            return found_paragraphs[:5]  # Ограничиваем 5 абзацами на файл
            
        except Exception as e:
            print(f"Ошибка поиска в файле {file_path}: {e}")
            return []